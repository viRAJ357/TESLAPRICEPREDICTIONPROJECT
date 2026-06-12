import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = docx.Document()

# Title
title = doc.add_heading('Tesla Price Prediction: 1.5 Hour Workshop Speech', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph('A Complete Beginner-Friendly Guide')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph("--------------------------------------------------------------------------------------------------\n")
doc.add_paragraph("Note to Speaker: This script is designed for a 90-minute (1.5-hour) presentation. It includes pauses, analogies, and crowd interactions. Speak slowly, naturally, and use hand gestures. Whenever you see [PAUSE], take a deep breath and look at the audience.\n")

def add_section(time_marker, heading_text, content_paragraphs):
    heading = doc.add_heading(f"{time_marker} - {heading_text}", level=1)
    for p in content_paragraphs:
        doc.add_paragraph(p)
    doc.add_paragraph("")

# Section 1
add_section("0:00 to 0:15", "Introduction & Icebreaker", [
    "Hello everyone! Welcome. Thank you all for joining today. [PAUSE]",
    "Let me ask you a very simple question. If you had a time machine, and you could go back to the year 2010, what is the one thing you would do? [PAUSE to let people answer or think]",
    "Many of us would say: 'I would buy Tesla stock!' Why? Because a few dollars invested back then would be worth a fortune today.",
    "But we don't have a time machine. What we do have is Artificial Intelligence. Today, we are going to look at a project that tries to act like a mini time machine. We have built an AI that looks at the past stock prices of Tesla and tries to predict its future price.",
    "I promise you, by the end of this 1.5 hours, you will understand exactly how we built this, even if you have never written a single line of code in your life. We will keep it simple, easy, and fun."
])

# Section 2
add_section("0:15 to 0:30", "Understanding the Data (The Fuel for our AI)", [
    "Alright, let's start with step one: The Data.",
    "If you want to prepare for an exam, what do you need? You need past exam papers, right? The same goes for Artificial Intelligence. To predict the future, our AI needs to study the past.",
    "In our project, we have a file named 'TSLA.csv'. Think of a CSV file as a simple Excel spreadsheet. Inside this spreadsheet, we have thousands of rows. Each row tells us: On this specific date, what was the price of Tesla?",
    "But we cannot just throw this data at the AI. If I give you a math book written in a language you don't understand, you won't learn anything. We have to prepare the data. We call this 'Data Preprocessing'.",
    "We use a tool called a 'MinMaxScaler'. Imagine you have prices ranging from $10 to $1,000. Big numbers can confuse the AI. So, we shrink all the numbers down to a scale between 0 and 1. Zero is the lowest price, and One is the highest. It makes the math super easy for our computer brain. [PAUSE]",
    "Next, we create 'sequences'. We tell the AI: 'Look at the last 60 days of prices. Now, guess the 61st day.' We do this over and over again, thousands of times. This is how the AI studies!"
])

# Section 3
add_section("0:30 to 0:55", "The Brain of the AI (LSTM & SimpleRNN)", [
    "Now we come to the most exciting part. The Brain. We are using Deep Learning. Deep Learning is just a fancy word for a computer system inspired by the human brain.",
    "We used two different types of brains in this project. The first one is called 'SimpleRNN' (Recurrent Neural Network).",
    "Imagine you are reading a book. You understand the current word based on the previous words you just read. You don't read every word in isolation. SimpleRNN does exactly this! It looks at today's price, but it remembers yesterday's price.",
    "But SimpleRNN has a problem. It has a bad memory. If you ask it about a price from 50 days ago, it forgets. [PAUSE]",
    "To fix this, we used a second, smarter brain called 'LSTM'. It stands for Long Short-Term Memory.",
    "LSTM is like a student who knows what to remember for the final exam and what to forget. It has special 'gates'. If a piece of information is useless, the forget gate throws it away. If it is important, it keeps it in its long-term memory.",
    "When we trained both these brains, we saw that the LSTM was much better at guessing Tesla's stock price because it could remember the long-term trends of the stock market."
])

# Section 4
add_section("0:55 to 1:15", "Building the Face (The Streamlit Web App)", [
    "We have a smart AI brain, but a brain inside a computer is boring. We needed a face. A beautiful interface that you, me, or a layman can use.",
    "To build this, we used a magical tool called 'Streamlit'. Streamlit allows us to build a website in just a few lines of Python code.",
    "But we didn't stop at making a basic website. We made it 'Dynamic, Scalable, and Smooth'.",
    "What does that mean?",
    "First, we used 'Plotly' for our charts. Unlike boring static pictures, Plotly charts are interactive. When you move your mouse over the chart, you can see the exact date and price popping up. You can zoom in and drag it around.",
    "Second, we added smooth CSS animations. When you open the app, the numbers and charts smoothly fade into the screen. When you hover your mouse over the predicted price, it beautifully lifts up with a red glowing shadow. It feels like a premium, expensive software.",
    "Third, we made it multifunctional. We added a 'Search by Date' feature. Let's say you want to know what the AI thinks the price will be exactly 5 days from now. You just select the date from a drop-down menu, and boom! The exact price is displayed. We also added a button so you can download the AI's predictions straight to your computer as an Excel file."
])

# Section 5
add_section("1:15 to 1:25", "Live Demo & Walkthrough", [
    "[Note to Speaker: At this point, open your laptop, run 'python -m streamlit run app.py' and share your screen.]",
    "Let me show you exactly what we built. Here on my screen is the Tesla Price Predictor.",
    "Look at this beautiful chart showing Tesla's historical data. As I move my mouse, you can see how interactive it is.",
    "On the left, I can choose my AI Brain: LSTM or SimpleRNN. Let's choose LSTM.",
    "Now I will tell it to forecast the next 10 days. I click 'Generate Forecast'.",
    "Watch the progress bar... The AI is thinking... and there it is! A beautiful red dashed line appears, showing the future.",
    "Down here, I can use the search tool. Let's pick next Wednesday... and it tells me the exact predicted price. And if I want to send this to my boss, I just click this 'Download Forecast' button.",
    "[PAUSE and let the audience appreciate the demo]"
])

# Section 6
add_section("1:25 to 1:30", "Conclusion & Q&A", [
    "To wrap things up: We took raw, confusing stock data. We cleaned it. We fed it to a deep learning AI called LSTM that learned the hidden patterns. And finally, we wrapped it all in a beautiful, smooth web application.",
    "Is this AI a crystal ball? No. The stock market is highly unpredictable. A tweet from Elon Musk can change the price in seconds! The AI doesn't read tweets; it only reads past prices.",
    "But what we have built is a powerful tool to understand trends and momentum.",
    "Thank you so much for your time and attention over the past hour and a half. I hope you found this exciting and easy to understand.",
    "I am now open to any questions you might have!"
])

doc.save('Project_Explanation_Speech.docx')
print('Speech DOCX generated successfully.')
