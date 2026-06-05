import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = docx.Document()

# Title
title = doc.add_heading('Tesla Stock Price Prediction using Deep Learning', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Subtitle
subtitle = doc.add_paragraph('Project Report')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_format = subtitle.runs[0].font
subtitle_format.size = Pt(14)
subtitle_format.bold = True

doc.add_paragraph('\nSubmitted To: Labmentix')
doc.add_paragraph('--------------------------------------------------\n')

# 1. Introduction
doc.add_heading('1. Introduction', level=1)
doc.add_paragraph("Stock market prediction is the act of trying to determine the future value of a company stock or other financial instrument traded on an exchange. The successful prediction of a stock's future price could yield significant profit. In this project, we aim to predict the future stock prices of Tesla Inc. (TSLA) using advanced Deep Learning techniques.")
doc.add_paragraph("Time-series forecasting, specifically using Recurrent Neural Networks (RNN) and Long Short-Term Memory (LSTM) networks, has proven highly effective in capturing the temporal dependencies and non-linear patterns inherent in stock market data.")

# 2. Objective
doc.add_heading('2. Project Objective', level=1)
doc.add_paragraph("The primary objectives of this project are:")
doc.add_paragraph("• To develop and train Deep Learning models (LSTM and SimpleRNN) to forecast Tesla's stock prices.", style='List Bullet')
doc.add_paragraph("• To build an interactive, dynamic, and scalable web application allowing users to visualize historical data and future predictions.", style='List Bullet')
doc.add_paragraph("• To provide a multi-functional user interface where users can search forecasts by specific dates and download prediction reports.", style='List Bullet')

# 3. Methodology & Implementation
doc.add_heading('3. Methodology & Implementation', level=1)
doc.add_heading('3.1 Data Collection & Preprocessing', level=2)
doc.add_paragraph("Historical stock data for Tesla (TSLA) was collected and preprocessed. The 'Adj Close' price was isolated as the target variable. The data was normalized using MinMaxScaler to scale the values between 0 and 1, which is crucial for the convergence of neural networks. We utilized a sequence length (time step) of 60 days to predict the next day's price.")

doc.add_heading('3.2 Deep Learning Architectures', level=2)
doc.add_paragraph("Two different recurrent architectures were implemented and compared:")
doc.add_paragraph("1. SimpleRNN: A standard Recurrent Neural Network designed to handle sequential data, though susceptible to the vanishing gradient problem.")
doc.add_paragraph("2. LSTM (Long Short-Term Memory): An advanced RNN architecture with internal gates (forget, input, output) specifically designed to remember long-term dependencies and overcome the vanishing gradient problem.")

doc.add_heading('3.3 Web Application Interface', level=2)
doc.add_paragraph("A dynamic web application was built using the Streamlit framework. Key features include:")
doc.add_paragraph("• Interactive Data Visualization: Integration of Plotly for smooth, scalable, and interactive charting (zoom, pan, hover).", style='List Bullet')
doc.add_paragraph("• Custom Configurations: Users can choose between LSTM and SimpleRNN models, and adjust the forecast horizon (up to 30 days).", style='List Bullet')
doc.add_paragraph("• Multi-functional Capabilities: A built-in feature to search for a specific forecasted date's price, along with the ability to export the forecast as a CSV file.", style='List Bullet')
doc.add_paragraph("• Real-time CSS Animations: A modern UI with hover transitions and metric box animations for a premium feel.", style='List Bullet')

# 4. Results & Conclusion
doc.add_heading('4. Results & Conclusion', level=1)
doc.add_paragraph("Both models successfully learned the underlying trends of the Tesla stock. The LSTM model generally outperformed the SimpleRNN by capturing the longer-term volatility and resisting sudden noise, reflecting a lower Mean Squared Error (MSE).")
doc.add_paragraph("The integration of the model into a fully-fledged, interactive dashboard provides a comprehensive tool for financial analysis and prediction.")

doc.save('Labmentix_Project_Report.docx')
print('Report generated successfully.')
